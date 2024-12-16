import os
import re
import operator
import toml

from tempfile import gettempdir
from pathlib import Path
from datetime import datetime, timedelta
from tzlocal import get_localzone
from pytz import utc as UTC
from subprocess import Popen, PIPE

from utils import app
from ivsdb import IVSdata, models
from vgosdb import VGOSdb

from pandas.tseries.holiday import USFederalHolidayCalendar

from docx import Document, shared
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document


class Report:

    not_defined = datetime(1970, 1, 1)

    def __init__(self, monthly=False, today=None, first_day=0):

        self.end, self.start, self.submitted, self.period = None, None, None, None
        self.template, self.path = None, None
        self.scheduled = []
        self.analyzed = {'standard': [], 'intensive': []}
        self.is_vgos = {'standard': 0, 'intensive': 0}

        today = datetime.strptime(today, '%Y-%m-%d') if today \
            else datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

        info = app.load_control_file(name=app.ControlFiles.NVIreport)[-1]

        self.monthly(info, today) if monthly else self.weekly(info, today, first_day)

        self.scheduling_comments = info['Scheduled']
        with open(Path(Path(app.args.config).parent, 'logger.toml')) as log_config:
            self.current_log = Path(toml.load(log_config)['handlers']['file']['filename'])

    # Set parameters for weekly report
    def weekly(self, info, today, first_day):
        if (days := today.weekday() - first_day + 1) < 0:
            days += 7
        self.end = today - timedelta(days=days)
        self.start = self.end - timedelta(days=6)
        self.submitted = self.start + timedelta(days=10) if first_day == 0 else today
        self.template = info['Templates']['weekly']

        if self.start.strftime('%Y') != self.end.strftime('%Y'):
            self.period = f'{self.start.strftime("%B %-d, %Y")} to {self.end.strftime("%B %-d, %Y")}'
        elif self.start.strftime('%B') != self.end.strftime('%B'):
            self.period = f'{self.start.strftime("%B %-d")} to {self.end.strftime("%B %-d, %Y")}'
        else:
            self.period = f'{self.start.strftime("%B %-d")} to {self.end.strftime("%-d, %Y")}'

        self.path = os.path.join(gettempdir(), self.submitted.strftime('%Y-%m-%d.docx'))

    # Set parameters for monthly report
    def monthly(self, info, today):
        self.end = today.replace(day=1) - timedelta(days=1)
        self.start = self.end.replace(day=1)
        self.end += timedelta(days=1)
        self.submitted = self.end + timedelta(days=10)
        self.template = info['Templates']['monthly']

        self.period = self.start.strftime('%B %Y')
        self.path = os.path.join(gettempdir(), self.start.strftime('%Y-%m.docx'))

    # Special function to decode timetag
    def decode_timetag(self, line):
        # Some time the datetime.strptime failed because of seconds = 60 or hours = 24.
        data = line.replace('TIMETAG', '').replace('UTC', '').strip()
        hour, minute, second = list(map(int, data[-8:].split(':')))
        seconds = second + minute * 60 + hour * 3600
        utc = datetime.strptime(data[:10], '%Y/%m/%d') + timedelta(seconds=seconds)
        return UTC.localize(utc)

    def is_auto(self, ses):
        if not ses.is_intensive:
            return False
        start, lines = ses.start.strftime('%Y-%m-%d.%h%m%s'), []
        for file in self.current_log.parent.iterdir():
            if file.suffix == '.gz' and file.stem > start:
                cmd = f"zgrep {ses.db_name} {file}"
                st_out, st_err = Popen(cmd, shell=True, stdin=PIPE, stdout=PIPE, stderr=PIPE).communicate()
                lines.extend(st_out.decode('utf-8').splitlines())
        cmd = f"grep {ses.db_name} {self.current_log}"
        st_out, st_err = Popen(cmd, shell=True, stdin=PIPE, stdout=PIPE, stderr=PIPE).communicate()
        lines.extend(st_out.decode('utf-8').splitlines())
        data = '\n'.join(lines)
        find = re.compile(r'aps .* failed').findall
        return all(('auto analyse' in data, 'nuSolve ok' in data, 'send auto processing email' in data,
                    not bool(find(data))))

    # Get analyzed time from history file
    def get_analyzed_time(self, vgosdb, nuSolve):
        path = os.path.join(vgosdb.folder, 'History', nuSolve['history'])
        timetags = []
        with open(path) as hist:
            for line in hist:
                if line.lstrip().startswith('TIMETAG'):
                    timetags.append(self.decode_timetag(line).astimezone(get_localzone()).replace(tzinfo=None))
        timetags.sort()
        return timetags[-1]

    # Get list of vgosdb
    def get_vgosdb(self, dbase):
        # Define time limits and pattern to find files
        start, end = self.start.timestamp(), self.end.timestamp()
        found, pattern = [], '*/*/*GSFC_kall.wrp'

        # Find all vgosdb wrappers that have been updated during that period
        root = app.VLBIfolders.vgosdb
        found = [file.parent.name.split('.')[0] for file in Path(root).glob(pattern)
                 if start <= file.stat().st_mtime <= end]

        # Get vgosdb files that have been analyzed during that period
        for db_name in set(found):
            if ses_id := dbase.get_db_session_code(db_name):
                if session := dbase.get_session(ses_id):
                    vgosdb = VGOSdb(session.db_folder)
                    downloaded = None
                    if wrapper := vgosdb.get_first_wrapper('GSFC'):
                        # Get time of vgosDbCalc (first action after download)
                        if calc := wrapper.processes.get('vgosDbCalc', None):
                            downloaded = calc['runtimetag'].astimezone(get_localzone()).replace(tzinfo=None)
                    if wrapper := vgosdb.get_last_wrapper('GSFC'):
                        if nuSolve := wrapper.processes.get('nuSolve', None):
                            analyzed = self.get_analyzed_time(vgosdb, nuSolve)
                            if not downloaded or downloaded > analyzed:
                                downloaded = self.get_downloaded(dbase, db_name, analyzed)
                            if self.start <= analyzed <= self.end:
                                record = (ses_id, session.start, downloaded, analyzed,
                                          self.latency(session, downloaded, analyzed),
                                          self.comments(session))
                                self.analyzed[session.type].append(record)
                                self.is_vgos[session.type] += wrapper.has_cal_cable()

    # Get time the file has been downloaded
    def get_downloaded(self, dbase, db_name, analyzed):
        pattern = f'{db_name}%'
        if not (lst := dbase.orm_ses.query(models.CorrFile).filter(models.CorrFile.code.like(pattern)).all()):
            return None
        lst.sort(key=operator.attrgetter('updated'), reverse=True)
        for info in lst:
            if analyzed >= info.updated:
                return info.updated
        lst.sort(key=operator.attrgetter('first'), reverse=True)
        for info in lst:
            if info.first != self.not_defined and analyzed >= info.first:
                return info.first
        return None

    # Get processing latency
    def latency(self, session, downloaded, processed):
        if not downloaded or downloaded > processed:
            return '??'
        if self.is_auto(session):
            return 'AUTO'
        dt = (processed - downloaded).total_seconds()
        hours, minutes = divmod(dt / 60, 60)
        if minutes > 59.5:
            hours += 1
            minutes = 0
        if hours > 23:
            days, hours = divmod(dt / 3600, 24)
            if hours > 23.5:
                days += 1
                return f'{int(days):d}day{"s" if int(days) > 1 else ""}'
            return f'{int(days):d}d {int(hours):02d}h'

        return '{:02d}:{:02d}'.format(int(hours), int(minutes))

    # Get comments from analysis report
    @staticmethod
    def comments(session):
        agency = 'IVS' if session.analysis_center.upper() == 'NASA' else 'NASA'
        comments = []
        if reports := sorted(Path(session.folder).glob(f'*{agency}-analysis-report*.txt')):
            with open(reports[-1]) as rep:
                for line in rep:
                    if line.startswith('Problems:'):
                        comments.append(re.sub('^- ', '', line.split('Problems:')[-1].replace('None.', '')).strip())
                    elif line.startswith(('Number of', 'Parameterization')):
                        break
                    elif comments:
                        comments.append(re.sub('^- ', '',line.strip()))
        # Clean comments
        return '\n'.join(list(filter(None, comments)))

    # Get list of upcoming schedules
    def get_scheduled(self, dbase):
        first, last = self.start, datetime.now() + timedelta(days=31)
        upcoming = [session for ses_id in dbase.get_sessions(first, last, ['standard', 'intensive'])
                    if (session := dbase.get_session(ses_id)) and session.operations_center.upper() == 'NASA']

        # Check if the schedule files have been downloaded on our servers
        for session in upcoming:
            codes = [f'{session.code}.skd', f'{session.code}.vex']
            if lst := dbase.orm_ses.query(models.RecentFile).filter(
                    models.RecentFile.code.in_(codes)).order_by(models.RecentFile.first.asc()).all():
                downloaded = lst[0].first  # Local time
                if self.start <= downloaded <= self.end:
                    days = int((session.start.date() - downloaded.date()).total_seconds() / 86400)
                    readiness = f'{days:d} day{"s" if days > 1 else ""}'
                    comments = self.scheduling_comments.get(session.code, '')
                    self.scheduled.append((session.code, session.start, downloaded, readiness, comments))

    # Create the docx report
    def create(self):

        # Function to iterate blocks and find Tables or Paragraphs
        def iter_blocks(parent):
            if isinstance(parent, _Document):
                element = parent.element.body
            elif isinstance(parent, Table._Cell):
                element = parent._tc
            else:
                raise ValueError("something's not right")

            for child in element.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        # Update cell and change color if warning
        def update_cell(cell, text, alignment, warning=False):
            cell.text = text
            cell.paragraphs[0].alignment = alignment
            if warning:
                cell.paragraphs[0].runs[0].font.color.rgb = shared.RGBColor(255, 0, 0)

        doc = Document(self.template)

        # Initiate keywords to replace in text
        keys = {'{DATE}': self.submitted.strftime('%Y-%m-%d'), '{PERIOD}': self.period}
        tables = []
        # Read all paragraphs to find tables and change keywords in text
        for block in iter_blocks(doc):
            if isinstance(block, Table):
                tables.append(block)
            elif isinstance(block, Paragraph):
                lines = block.runs
                for i in range(len(lines)):
                    for key, word in keys.items():
                        if key in lines[i].text:
                            lines[i].text = lines[i].text.replace(key, word)

        # Get US holidays and define some formatting functions
        holidays = USFederalHolidayCalendar().holidays(start=str(self.start), end=str(self.end))
        hm = lambda x: x.strftime('%b%d %H:%M') if x else 'N/A'
        ymd = lambda x: x.strftime('%Y-%m-%d')
        warning = lambda d: ymd(d) in holidays or d.weekday() > 4 or not (9 <= d.hour < 18)

        # Populate tables for Standard 24Heure and Intensive sessions
        for name, index in {'standard': -3, 'intensive': -2}.items():
            table = tables[index]
            # Get alignment from header row
            alignments = [cell.paragraphs[0].alignment for cell in table.rows[0].cells]

            sessions = sorted(self.analyzed[name], key=operator.itemgetter(3))
            for (ses_id, start, downloaded, analyzed, latency, comments) in sessions:
                if not downloaded or not analyzed:
                    continue
                row = table.add_row()
                try:
                    update_cell(row.cells[0], ses_id, alignments[0])
                    update_cell(row.cells[1], ymd(start), alignments[1])
                    update_cell(row.cells[2], hm(downloaded), alignments[2], warning(downloaded))
                    update_cell(row.cells[3], hm(analyzed), alignments[3])
                    update_cell(row.cells[4], latency, alignments[4])
                    update_cell(row.cells[5], comments, alignments[5])
                except Exception as err:
                    print(f'Problem with {ses_id}: {str(err)}')
                    print(ses_id, start, downloaded, analyzed, latency, comments)


        # Populate table for scheduled sessions
        table = tables[-1]
        # Get alignment from header row
        alignments = [cell.paragraphs[0].alignment for cell in table.rows[0].cells]

        sessions = sorted(self.scheduled, key=operator.itemgetter(1))
        for (ses_id, start, available, readiness, comments) in sessions:
            row = table.add_row()
            update_cell(row.cells[0], ses_id, alignments[0])
            update_cell(row.cells[1], ymd(start), alignments[1])
            update_cell(row.cells[2], ymd(available), alignments[2])
            update_cell(row.cells[3], readiness, alignments[3])
            update_cell(row.cells[4], comments, alignments[4])

        # Save document
        doc.save(self.path)


class AnalyzedReport(Report):

    def __init__(self, year):
        super().__init__()
        self.start, self.end = datetime(year, 1, 1), datetime(year + 1, 1, 1)


def report_analyzed(dbase, year):
    # Define time limits and pattern to find files
    start, end = datetime(year, 1, 1).timestamp(), datetime(year+1, 1, 1).timestamp()
    found, pattern = [], '*/*/*GSFC_kall.wrp'

    # Find all vgosdb wrappers that have been updated during that period
    root = app.VLBIfolders.vgosdb
    found = [file.parent.name.split('.')[0] for file in Path(root).glob(pattern)
             if start <= file.stat().st_mtime <= end]

    # Get vgosdb files that have been analyzed during that period
    index = 0
    stats = dict(standard=0, intensive=0)
    nbr_per_year = dict(standard={}, intensive={})
    for key in nbr_per_year:
        nbr_per_year[key] = {str(y): 0 for y in range(1980, 2023)}
    for db_name in set(found):
        if ses_id := dbase.get_db_session_code(db_name):
            if session := dbase.get_session(ses_id):
                stats[session.type] += 1
                nbr_per_year[session.type][session.year] += 1
                index += 1
                print(index, session.type, session)
    print(stats)
    print(nbr_per_year)


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser( description='Logger for ADAP software.' )
    parser.add_argument('-c', '--config', help='config file', required=True)
    parser.add_argument('-d', '--db', help='database name', default='ivscc', required=False)
    parser.add_argument('-a', '--analyzed', help='analyzed sessions for year', type=int, default=0, required=False)
    parser.add_argument('-m', '--monthly', help='monthly report', action='store_true')
    parser.add_argument('-t', '--today', help='change today', required=False)
    parser.add_argument('-f', '--first_day', help='first week day of report (0 is Monday)',
                        type=int, default=0, required=False)

    args = app.init(parser.parse_args())

    db_url = app.load_control_file(name=app.ControlFiles.Database)[-1]['Credentials'][app.args.db]

    with IVSdata(db_url) as dbase:
        if args.analyzed:
            report = AnalyzedReport(args.analyzed)
            report.get_vgosdb(dbase)
            for key, val in report.analyzed.items():
                print(key, len(val), report.is_vgos[key])
        else:
            report = Report(args.monthly, args.today, args.first_day)
            report.get_vgosdb(dbase)
            report.get_scheduled(dbase)
            report.create()
            print(f'Report: {report.path}')


